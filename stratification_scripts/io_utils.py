"""
I/O utilities for the Federal Register Study pipeline.

This module provides:
- UTF-8 text sanitization for API responses
- PDF text extraction
- DOCX text extraction

Example:
    >>> from stratification_scripts.io_utils import sanitize_utf8, extract_pdf_text, extract_docx_text
    >>>
    >>> # Clean text with invalid UTF-8
    >>> clean_text = sanitize_utf8(dirty_text)
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any, Optional

from stratification_scripts.logging_utils import get_logger

logger = get_logger(__name__)


def sanitize_utf8(text: Any) -> Optional[str]:
    """
    Sanitize text to remove invalid UTF-8 characters.

    Removes surrogate characters (U+D800 to U+DFFF) that can't be encoded
    to valid UTF-8, which can cause issues when writing to CSV or databases.

    Args:
        text: Input text (str, bytes, or None)

    Returns:
        Sanitized UTF-8 string or None if input is None/empty.

    Example:
        >>> sanitize_utf8("Hello \\ud83d\\ude00 World")  # Valid emoji
        'Hello World'
        >>> sanitize_utf8(None)
        None
    """
    if text is None:
        return None

    if isinstance(text, bytes):
        try:
            text = text.decode('utf-8', errors='replace')
        except Exception:
            return None

    if not isinstance(text, str):
        text = str(text)

    if not text:
        return None

    # Remove surrogate characters that can't be encoded to UTF-8
    # Surrogates are in the range U+D800 to U+DFFF
    sanitized_chars = []
    for char in text:
        code_point = ord(char)
        # Skip surrogate characters
        if 0xD800 <= code_point <= 0xDFFF:
            continue
        sanitized_chars.append(char)

    sanitized = ''.join(sanitized_chars)

    if not sanitized:
        return None

    # Final check: ensure it can be encoded to UTF-8
    try:
        sanitized.encode('utf-8')
        return sanitized
    except (UnicodeEncodeError, UnicodeError):
        # Last resort: use error handling
        try:
            return sanitized.encode('utf-8', errors='replace').decode('utf-8', errors='replace') or None
        except Exception:
            return None


def extract_pdf_text(
    pdf_bytes: bytes,
    max_pages: int = 2,
) -> Optional[str]:
    """
    Extract text from PDF bytes.

    Tries pypdf first, falls back to PyMuPDF if that fails.

    Args:
        pdf_bytes: Raw PDF file content
        max_pages: Maximum number of pages to extract (default: 2)

    Returns:
        Extracted text or None if extraction failed.
    """
    if not pdf_bytes:
        return None

    # Verify it's a PDF
    if not pdf_bytes.startswith(b'%PDF'):
        if b'<!DOCTYPE html>' in pdf_bytes[:100] or b'<html' in pdf_bytes[:100]:
            logger.debug("Got HTML instead of PDF")
            return None
        logger.debug(f"Not a PDF file (starts with {pdf_bytes[:20]})")
        return None

    # Try pypdf first (pure python, often reliable)
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(pdf_bytes))

        text_parts = []
        for page_num in range(min(max_pages, len(reader.pages))):
            page = reader.pages[page_num]
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        text = "\n\n".join(text_parts).strip()
        return sanitize_utf8(text) if text else None

    except Exception as pypdf_error:
        # Fall back to PyMuPDF
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=io.BytesIO(pdf_bytes), filetype="pdf")

            text_parts = []
            for page_num in range(min(max_pages, len(doc))):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text:
                    text_parts.append(page_text)

            doc.close()

            text = "\n\n".join(text_parts).strip()
            return sanitize_utf8(text) if text else None

        except Exception as pymupdf_error:
            logger.warning(
                f"PDF extraction failed: pypdf={pypdf_error}, "
                f"pymupdf={pymupdf_error}"
            )
            return None


def extract_docx_text(
    docx_bytes: bytes,
    max_pages: int = 2,
) -> Optional[str]:
    """
    Extract text from DOCX bytes.

    Uses python-docx to read the document and extract paragraph text.
    Approximates page limits by capping at ~3000 chars (roughly 2 pages).

    Args:
        docx_bytes: Raw DOCX file content
        max_pages: Maximum pages to approximate (default: 2)

    Returns:
        Extracted text or None if extraction failed.
    """
    if not docx_bytes:
        return None

    # DOCX files are ZIP archives starting with PK signature
    if not docx_bytes[:2] == b'PK':
        logger.warning(f"Not a DOCX file (starts with {docx_bytes[:20]})")
        return None

    try:
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        text_parts = []
        char_count = 0
        # ~1500 chars per page approximation
        char_limit = max_pages * 1500

        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                text_parts.append(para_text)
                char_count += len(para_text)
                if char_count >= char_limit:
                    break

        # Also extract text from tables (common in regulatory comments)
        if char_count < char_limit:
            for table in doc.tables:
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        text_parts.append(" | ".join(row_texts))
                        char_count += sum(len(t) for t in row_texts)
                        if char_count >= char_limit:
                            break
                if char_count >= char_limit:
                    break

        text = "\n\n".join(text_parts).strip()
        return sanitize_utf8(text) if text else None

    except Exception as e:
        logger.warning(f"DOCX extraction failed: {e}")
        return None
